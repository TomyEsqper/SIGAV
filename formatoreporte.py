from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import os
import sys
import json

MESES_ES = {
    1: 'ENERO', 2: 'FEBRERO', 3: 'MARZO', 4: 'ABRIL', 5: 'MAYO',
    6: 'JUNIO', 7: 'JULIO', 8: 'AGOSTO', 9: 'SEPTIEMBRE', 10: 'OCTUBRE',
    11: 'NOVIEMBRE', 12: 'DICIEMBRE'
}


def mes_es(m):
    try:
        return MESES_ES[int(m)]
    except Exception:
        return ''


def parse_fecha(value):
    if isinstance(value, datetime):
        return value
    if isinstance(value, str):
        for fmt in ('%Y-%m-%d', '%Y-%m-%dT%H:%M:%S', '%Y-%m-%dT%H:%M:%S.%f', '%d/%m/%Y', '%d-%m-%Y'):
            try:
                return datetime.strptime(value, fmt)
            except Exception:
                pass
    return datetime.now()


def set_bold_line(paragraph, label, value):
    paragraph.text = ''
    r1 = paragraph.add_run(label)
    r1.bold = True
    r2 = paragraph.add_run(value)
    r2.bold = True


def generar_informe_camaras(datos):
    base_dir = os.path.dirname(__file__)
    template_path = os.path.join(base_dir, 'FORMATO_DE_INFORME[1].docx')
    doc = Document(template_path)

    # Encabezado fecha (p0)
    fa_dt = parse_fecha(datos.get('fecha_actual'))
    fecha_enc = f"{fa_dt.day:02d} de {mes_es(fa_dt.month)} del {fa_dt.year}"
    p_enc = doc.paragraphs[0]
    p_enc.text = f"Ibagué, {fecha_enc}"
    p_enc.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Días relacionados (p10)
    p_dias = doc.paragraphs[10]
    set_bold_line(p_dias, 'DÍAS RELACIONADOS: ', str(datos.get('dias_revisados', '')).upper())

    # Vehículo (p11) -> "  VEHÍCULO: CO-0001  TGVXXX"
    p_veh = doc.paragraphs[11]
    codigo = str(datos.get('vehiculo_codigo', '')).upper()
    placa = str(datos.get('vehiculo_placa', '')).upper()
    p_veh.text = ''
    r_lbl = p_veh.add_run('  VEHÍCULO: ')
    r_lbl.bold = True
    r_val = p_veh.add_run(f"{codigo}  {placa}")
    r_val.bold = True

    # Conductor (p12)
    p_cond = doc.paragraphs[12]
    set_bold_line(p_cond, 'CONDUCTOR: ', str(datos.get('conductor', '')))

    # Número de informe (p13)
    p_num = doc.paragraphs[13]
    set_bold_line(p_num, 'NUMERO DE INFORME: ', str(datos.get('numero_informe', '')))

    # Cuerpo (p15)
    fecha_rev = str(datos.get('fecha_revision', ''))
    ruta = str(datos.get('numero_ruta', ''))
    p_body = doc.paragraphs[15]
    p_body.text = (
        f"El día {fecha_rev} del presente año, el vehículo que cubría la ruta No.{ruta} "
        f"presentó el siguiente registro de pasajeros que ingresaron por la puerta trasera"
    )

    # Tabla de registros (tabla 0)
    tbl = doc.tables[0]
    hdr = tbl.rows[0].cells
    hdr[0].text = 'GRABACIÓN'
    hdr[1].text = 'HORA'
    hdr[2].text = 'No. PASAJEROS'
    for c in hdr:
        for par in c.paragraphs:
            for run in par.runs:
                run.bold = True
        for par in c.paragraphs:
            par.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Limpiar filas existentes (dejar solo encabezado)
    while len(tbl.rows) > 1:
        tbl._tbl.remove(tbl.rows[1]._tr)

    for reg in datos.get('registros', []):
        row = tbl.add_row().cells
        row[0].text = str(reg.get('grabacion', ''))
        row[1].text = str(reg.get('hora', ''))
        row[2].text = str(reg.get('pasajeros', 0))
        for cell in row:
            for par in cell.paragraphs:
                par.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Fila TOTAL
    total = int(datos.get('total_pasajeros', 0))
    total_cells = tbl.add_row().cells
    total_cells[0].text = 'TOTAL'
    total_cells[0].merge(total_cells[1])
    total_cells[2].text = str(total)
    for par in total_cells[0].paragraphs:
        for run in par.runs:
            run.bold = True
        par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for par in total_cells[2].paragraphs:
        par.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Resumen (p18)
    p_res = doc.paragraphs[18]
    p_res.text = (
        f"El total de pasajeros ingresados por la puerta trasera el día {fecha_rev} del presente año, "
        f"fue de {total} pasajero{'s' if total != 1 else ''}."
    )

    nombre_archivo = datos.get('nombre_archivo') or os.path.join(base_dir, 'Informe_Revision_Camaras.docx')
    doc.save(nombre_archivo)
    return nombre_archivo


if __name__ == '__main__':
    if len(sys.argv) < 2:
        print('Uso: python formatoreporte.py <ruta_json>')
        raise SystemExit(1)
    input_path = sys.argv[1]
    if not os.path.isfile(input_path):
        print('Archivo JSON no encontrado:', input_path)
        raise SystemExit(1)
    with open(input_path, 'r', encoding='utf-8') as f:
        datos = json.load(f)
    ruta = generar_informe_camaras(datos)
    print(ruta)